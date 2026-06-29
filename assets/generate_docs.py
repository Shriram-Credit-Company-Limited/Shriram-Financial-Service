"""
generate_docs.py
Generates Shriram_Product_Pages_Content.docx from hardcoded product page data.

Run:
    pip install python-docx
    python generate_docs.py
"""

import sys
import subprocess

# ── Auto-install python-docx if missing ─────────────────────────────────────
try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ──────────────────────────────────────────────────────────────────

def add_page_break(doc):
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_table(doc, headers, rows):
    """Add a styled Word table."""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True

    # Data rows
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = str(val)

    doc.add_paragraph()  # spacing after table


def faq_section(doc, faqs):
    h2(doc, "Frequently Asked Questions")
    for item in faqs:
        h3(doc, item["q"])
        body(doc, item["a"])


def steps_section(doc, section_title, steps, eyebrow=None):
    h2(doc, section_title)
    for s in steps:
        h3(doc, f"{s['n']}. {s['title']}")
        body(doc, s["desc"])


def risks_section(doc, risks):
    h2(doc, "Risks")
    for r in risks:
        h3(doc, r["label"])
        body(doc, r["desc"])


def cta_section(doc, heading, paragraph):
    h2(doc, heading)
    body(doc, paragraph)


# ═══════════════════════════════════════════════════════════════════════════════
# PRODUCT PAGE DATA
# ═══════════════════════════════════════════════════════════════════════════════

PAGES = []

# ── 1. IPO ───────────────────────────────────────────────────────────────────
PAGES.append({
    "title": "Apply IPO Online India — ASBA UPI Upcoming IPOs 2025 | Shriram",
    "description": "Apply for upcoming IPOs online via ASBA or UPI on Shriram Financial Services. Track allotment, get Demat credit T+5/T+6, and zero brokerage on application. Start today.",
    "hero": {
        "eyebrow": "Initial Public Offering",
        "h1": "Apply for IPOs Before They Hit the Market",
        "subtitle": "Access all upcoming Mainboard and SME IPOs through the Shriram platform. Apply via UPI or ASBA, track allotment, and receive shares directly to your Demat.",
        "badges": ["Zero brokerage on application", "UPI & ASBA support", "Real-time allotment tracking", "Demat credit T+5/T+6"],
    },
    "sections": [
        {
            "eyebrow": "Types",
            "h2": "Types of IPO & Offer for Sale",
            "body": "Not all public offerings are the same. Understanding the type helps you evaluate risk, pricing, and post-listing behaviour.",
            "cards": [
                {"title": "Mainboard IPO", "tag": "Large Cap", "desc": "Large companies (Rs.10 crore+ paid-up capital) listing on NSE/BSE under full SEBI disclosure norms, typically well-covered by analysts."},
                {"title": "SME IPO", "tag": "High Growth", "desc": "SMEs with Rs.1-25 crore paid-up capital list on NSE Emerge or BSE SME. Higher growth potential - requires thorough due diligence."},
                {"title": "Offer for Sale (OFS)", "tag": "Secondary", "desc": "Promoters or PE investors sell existing stakes - no new capital raised by the company. Shares transfer directly from seller to buyer."},
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Apply for an IPO",
            "body": None,
            "steps": [
                {"n": "01", "title": "Check Upcoming IPOs", "desc": "Browse all live and upcoming IPOs on the Shriram platform. Read the DRHP, review financials, and check GMP and subscription trends."},
                {"n": "02", "title": "Enter Bid Details", "desc": "Choose the price band, lot size, and bid quantity. You can bid at the cut-off price or a specific price within the band."},
                {"n": "03", "title": "ASBA / UPI Mandate", "desc": "Funds are blocked in your bank account via ASBA or UPI mandate - no outflow until allotment. Blocked amounts earn interest."},
                {"n": "04", "title": "Allotment & Listing", "desc": "Allotment is done by lottery (if oversubscribed). Shares are credited to your Demat before listing day. Non-allotment amounts are unblocked immediately."},
            ],
        },
        {
            "eyebrow": "Timeline",
            "h2": "IPO Timeline - From Application to Listing",
            "body": None,
            "timeline": [
                {"label": "IPO Open Date", "desc": "Issue opens for bidding - typically 3 days. Price band and lot size are fixed."},
                {"label": "IPO Close Date", "desc": "Last day to place bids. Revisions or cancellations must be done before close time (typically 5 PM)."},
                {"label": "Basis of Allotment", "desc": "Published 6 days after close. Check allotment status on BSE/NSE or Shriram platform."},
                {"label": "Refund / Unblocking", "desc": "Non-allotted UPI mandates are released within T+1 day of allotment. ASBA amounts are also unblocked."},
                {"label": "Demat Credit", "desc": "Shares credited to your Demat account one day before listing, typically T+5 or T+6 from close date."},
                {"label": "Listing Day", "desc": "Stock begins trading on NSE/BSE. You can hold for the long term or book profits on listing day."},
            ],
        },
        {
            "eyebrow": "Know Before You Invest",
            "h2": "Risks of IPO Investing",
            "body": None,
            "risks": [
                {"label": "Overvaluation Risk", "desc": "IPO prices are set by investment bankers favouring the issuer. Some IPOs list at or below issue price - always analyse fundamentals before applying."},
                {"label": "Non-Allotment Risk", "desc": "Oversubscribed IPOs allocate by lottery in the retail category. You may bid and receive nothing - your funds are simply unblocked."},
                {"label": "Listing Volatility", "desc": "Stocks can swing wildly on listing day - listing high then crashing, or opening below issue price. Market sentiment dominates day-one price action."},
                {"label": "Lock-In Risk", "desc": "Anchor and promoter shares are locked for 30-90 days. Selling pressure after lock-in expiry can sharply depress the post-listing price."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is ASBA in an IPO application?", "a": "ASBA (Application Supported by Blocked Amount) blocks your bid amount in your bank account instead of transferring it. Interest keeps accruing; the amount is debited only on allotment and immediately unblocked if you don't get shares."},
        {"q": "Can I apply for IPO through Shriram Financial Services?", "a": "Yes. Apply via the Antara app or web portal using UPI or ASBA. View upcoming IPOs, track allotment status, and receive Demat credit - all in one place."},
        {"q": "What is the minimum lot size for IPO applications?", "a": "SEBI sets 1 lot as the retail minimum - typically Rs.10,000-Rs.15,000 per lot (10-200 shares). HNI/Non-Institutional Investors must bid a minimum of Rs.2 lakh."},
        {"q": "Is there a GMP (Grey Market Premium) and should I trust it?", "a": "GMP reflects unofficial pre-listing demand but is unregulated and unreliable. Use it as one data point only - always analyse the company's fundamentals before applying."},
        {"q": "How many IPO applications can I submit?", "a": "One application per PAN card in the Retail category. Duplicate PAN applications are rejected entirely. Family members with separate PANs can each apply independently."},
    ],
    "cta": {
        "h2": "Don't Miss the Next Big IPO",
        "body": "Open your Demat account and be ready for every upcoming IPO. Zero commission. Instant UPI application. Fast allotment tracking.",
    },
})

# ── 2. Derivatives ───────────────────────────────────────────────────────────
PAGES.append({
    "title": "F&O Trading India - Futures Options NSE | Shriram Financial",
    "description": "Trade futures and options on NSE with expert advisory, margin calculators, and real-time analytics. SEBI-registered. Flat Rs.20/order. Activate F&O segment today.",
    "hero": {
        "eyebrow": "Derivatives",
        "h1": "Trade Smarter with Futures & Options",
        "subtitle": "Access NSE and BSE derivative markets with advanced analytics, SEBI-registered research, and dedicated risk management support from Shriram Financial Services.",
        "badges": ["SEBI-registered", "Flat Rs.20/order", "Real-time analytics", "Expert RM support"],
    },
    "sections": [
        {
            "eyebrow": "The Basics",
            "h2": "What Are Derivatives",
            "body": "A derivative is a financial contract whose value is derived from an underlying asset - such as stocks, indices, commodities, or currencies. In India, derivatives are traded on NSE and BSE and are regulated by SEBI.",
            "cards": [
                {"title": "Futures Contracts", "desc": "A binding agreement to buy or sell an asset at a set price on a future date. Both parties are obligated to fulfil the contract at expiry."},
                {"title": "Options Contracts", "desc": "The right - not the obligation - to buy or sell at a strike price before expiry. Buyers risk only the premium paid; upside is unlimited for calls."},
            ],
        },
        {
            "eyebrow": "Requirements",
            "h2": "Eligibility & Documents Required",
            "body": None,
            "eligibility": [
                "Indian resident individual or entity (HUF, firm, company)",
                "Valid PAN Card and active trading + Demat account",
                "F&O segment activation on your trading account (SEBI mandate)",
                "Bank account with sufficient margin balance",
                "For NRI: RBI-approved broker and NRE/NRO bank account",
            ],
            "documents": [
                "PAN Card (mandatory for all account types)",
                "Address Proof - Aadhaar / Passport / Utility Bill",
                "Bank Account Proof - cancelled cheque or statement",
                "Income proof - 6-month bank statement or ITR (for F&O segment activation)",
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Start Trading F&O",
            "body": None,
            "steps": [
                {"n": "01", "title": "Open & Activate Trading Account", "desc": "Open a Demat + trading account and request F&O segment activation. Shriram enables this 100% digitally."},
                {"n": "02", "title": "Add Margin to Your Account", "desc": "Transfer funds to meet the SPAN + Exposure margin requirement. F&O requires upfront margin collateral."},
                {"n": "03", "title": "Select Underlying & Expiry", "desc": "Choose your index (NIFTY, BANKNIFTY) or stock, select the expiry date (weekly or monthly), and strike price."},
                {"n": "04", "title": "Place Your F&O Order", "desc": "Buy or write (sell) futures or options contracts through our platform and monitor P&L in real time."},
            ],
        },
        {
            "eyebrow": "Costs",
            "h2": "Brokerage & Charges",
            "body": None,
            "table": {
                "headers": ["Charge Type", "Futures", "Options"],
                "rows": [
                    ["Brokerage", "Rs.20 per executed order", "Rs.20 per executed order"],
                    ["STT", "0.0125% on sell side (contracts)", "0.125% on sell side (premium)"],
                    ["GST", "18% on brokerage + charges", "18% on brokerage + charges"],
                    ["Exchange Charges", "0.00235% (NSE)", "0.053% (NSE)"],
                    ["SEBI Turnover Fee", "Rs.10 per crore", "Rs.10 per crore"],
                ],
            },
        },
        {
            "eyebrow": "Strategy",
            "h2": "Popular F&O Strategies for Indian Traders",
            "body": None,
            "cards": [
                {"title": "Long Call", "tag": "Bullish", "desc": "Buy a call option when you expect the price to rise. Maximum loss is limited to the premium paid."},
                {"title": "Bull Put Spread", "tag": "Bullish", "desc": "Sell a higher-strike put, buy a lower-strike put. Profit if the underlying stays above the sold strike."},
                {"title": "Long Put", "tag": "Bearish", "desc": "Buy a put option when you expect the price to fall. Ideal for hedging an existing long position."},
                {"title": "Iron Condor", "tag": "Neutral", "desc": "Sell an OTM call spread and OTM put spread simultaneously. Profits in a low-volatility, range-bound market."},
            ],
        },
        {
            "eyebrow": "Know Before You Trade",
            "h2": "Risks of Trading Derivatives",
            "body": None,
            "risks": [
                {"label": "Leverage Risk", "desc": "F&O positions are leveraged - a small adverse move in the underlying can result in losses far exceeding the initial margin deposited."},
                {"label": "Expiry Risk", "desc": "Options lose time value rapidly as they approach expiry (theta decay). Positions left unmanaged near expiry can expire worthless."},
                {"label": "Volatility Risk", "desc": "Sharp, unexpected volatility events (gap openings, news events) can trigger margin calls or cause large overnight losses in open futures positions."},
                {"label": "Liquidity Risk", "desc": "Out-of-the-money contracts, particularly in individual stocks, may have very low open interest - making it difficult to exit at a fair price."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is the difference between Futures and Options?", "a": "Futures obligate both buyer and seller to transact at expiry. Options give the buyer the right - not the obligation - to buy (call) or sell (put). Options buyers risk only the premium; futures traders risk margin losses."},
        {"q": "How much capital do I need to start F&O trading?", "a": "Buying a NIFTY/BANKNIFTY options lot needs Rs.5,000-Rs.20,000 in premium. Futures require Rs.50,000-Rs.2,00,000 in SPAN + Exposure margin per lot, depending on the underlying."},
        {"q": "Is F&O income taxed differently from equity?", "a": "Yes - F&O P&L is business income (non-speculative), not capital gains. File ITR-3, offset losses against other business income, and carry forward losses for 8 years."},
        {"q": "What is 'writing' an option?", "a": "Writing (selling) an option earns you the premium upfront but creates an obligation to fulfil the contract if exercised. Writers profit when the option expires worthless; naked call writers face unlimited downside."},
        {"q": "How does margin work in F&O?", "a": "SEBI requires SPAN + Exposure margin upfront. SPAN covers worst-case one-day losses; Exposure adds a buffer. Fall below the required margin and your broker may square off positions."},
        {"q": "Can I trade F&O on any stock?", "a": "No. SEBI approves roughly 180+ stocks for F&O trading based on liquidity, market cap, and volatility. Index derivatives (NIFTY, BANKNIFTY, FINNIFTY) are the most liquid instruments available."},
    ],
    "cta": {
        "h2": "Start Trading Futures & Options Today",
        "body": "Activate F&O on your Shriram account in minutes. Access research-backed recommendations and trade with confidence.",
    },
})

# ── 3. ETF ───────────────────────────────────────────────────────────────────
PAGES.append({
    "title": "ETF Investment India - Buy Index Gold ETF NSE BSE | Shriram",
    "description": "Buy Index ETFs, Gold ETFs, and international ETFs on NSE & BSE. Ultra-low expense ratios, real-time pricing, and flat Rs.20 brokerage. Start ETF investing today.",
    "hero": {
        "eyebrow": "Exchange Traded Funds",
        "h1": "Trade ETFs - Diversify Your Portfolio Instantly",
        "subtitle": "Access Index ETFs, Gold ETFs, Sector ETFs, Debt ETFs, and International ETFs on NSE and BSE. Real-time pricing, ultra-low expense ratios, and the simplicity of stock trading.",
        "badges": ["Real-time trading", "Lowest expense ratios", "NSE & BSE listed", "All categories covered"],
    },
    "sections": [
        {
            "eyebrow": "Types",
            "h2": "Types of ETFs Available in India",
            "body": "The Indian ETF market has grown to Rs.7+ lakh crore in AUM. Here are the major categories you can access through Shriram.",
            "cards": [
                {"title": "Index ETFs", "desc": "Track benchmark indices like NIFTY 50, SENSEX, NIFTY Next 50. Low cost, broad diversification, and zero manager risk."},
                {"title": "Gold ETFs", "desc": "Invest in physical gold stored in vaults. Each unit = 1 gram of gold. No storage hassle, no purity risk, tradeable like stocks."},
                {"title": "Sectoral ETFs", "desc": "Track sector-specific indices: Bank Nifty, NIFTY IT, NIFTY Pharma, NIFTY FMCG. Concentrated exposure to a single industry."},
                {"title": "Debt ETFs", "desc": "Track fixed-income indices (government securities, corporate bonds). Lower volatility than equity ETFs, better liquidity than bond funds."},
                {"title": "International ETFs", "desc": "Exposure to global markets (US S&P 500, Nasdaq 100, MSCI World) through India-listed ETFs. Currency diversification built in."},
                {"title": "Smart Beta ETFs", "desc": "Factor-based strategies (Momentum, Low Volatility, Quality, Equal Weight) that go beyond simple index replication."},
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Start Investing in ETFs",
            "body": None,
            "steps": [
                {"n": "01", "title": "Open Demat & Trading Account", "desc": "ETFs are bought and sold like stocks on NSE/BSE. You need an active Demat + trading account with Shriram."},
                {"n": "02", "title": "Research ETFs", "desc": "Compare ETFs by expense ratio (TER), tracking error, liquidity (daily volumes), and underlying index. Lower tracking error = better replication."},
                {"n": "03", "title": "Place a Buy Order", "desc": "Place a market or limit order for the ETF on NSE/BSE during market hours (9:15 AM - 3:30 PM). ETFs settle T+1."},
                {"n": "04", "title": "Monitor & Rebalance", "desc": "Track NAV vs market price. Set up price alerts. Rebalance your ETF portfolio periodically to maintain target asset allocation."},
            ],
        },
        {
            "eyebrow": "Comparison",
            "h2": "ETF vs Mutual Funds",
            "body": None,
            "table": {
                "headers": ["Feature", "ETF", "Index Mutual Fund"],
                "rows": [
                    ["Trading", "Real-time throughout the day (like stocks)", "Once per day at end-of-day NAV"],
                    ["Minimum Investment", "Price of 1 unit (as low as Rs.30)", "Rs.500 (SIP) or Rs.1,000 (lump sum)"],
                    ["Expense Ratio", "0.04% to 0.40% (very low)", "0.5% to 2.5% (higher for active funds)"],
                    ["Manager Risk", "None (passive, index-tracking)", "High for active funds; none for index funds"],
                    ["Dividends", "Paid directly to investor or reinvested", "Dividend option or growth option"],
                    ["Tax Efficiency", "More tax-efficient (no internal churn)", "Less efficient for active funds (internal trading)"],
                    ["Demat Account", "Required", "Not required (can invest via AMC directly)"],
                ],
            },
        },
        {
            "eyebrow": "Know Before You Trade",
            "h2": "Risks of ETF Investing",
            "body": None,
            "risks": [
                {"label": "Tracking Error", "desc": "Small gaps between the ETF's return and the index it tracks can erode performance over time. Lower tracking error = better index replication."},
                {"label": "Liquidity Risk", "desc": "Sectoral and debt ETFs can have thin trading volumes and wide bid-ask spreads, making large trades expensive. Stick to high-volume ETFs for easier entry and exit."},
                {"label": "Market Risk", "desc": "Equity ETFs fall when the market falls - there is no active downside buffer. Use debt or hybrid ETFs to balance overall portfolio risk."},
                {"label": "Currency Risk", "desc": "International ETFs embed currency risk. If the INR appreciates against USD, your INR-denominated returns from a US ETF drop even if the underlying index rose."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is the difference between ETF and index mutual fund?", "a": "ETFs trade intraday on exchanges like stocks; index mutual funds can only be bought/sold once per day at NAV. ETFs have lower expense ratios but require a Demat account and incur brokerage. Index funds are simpler for SIP-based investors."},
        {"q": "Is a Demat account required to buy ETFs?", "a": "Yes. ETFs trade on NSE and BSE just like shares - you need an active Demat + trading account. Open one with Shriram and access all ETF categories alongside your equity portfolio."},
        {"q": "What is tracking error in ETFs?", "a": "Tracking error is the gap between the ETF's return and its benchmark index. Lower is better - the best NIFTY 50 ETFs maintain tracking errors below 0.05% per year."},
        {"q": "Can I start a SIP in an ETF?", "a": "ETFs don't support traditional auto-debit SIPs like mutual funds. You can manually buy units monthly, or use Shriram's systematic ETF purchase features where available."},
        {"q": "How are ETF dividends handled?", "a": "Most NIFTY 50 and Gold ETFs in India reinvest dividends into the NAV (growth option). Some ETFs pay dividends directly to investors when underlying holdings declare them."},
    ],
    "cta": {
        "h2": "Start Building Your ETF Portfolio",
        "body": "Access all ETF categories on NSE & BSE with flat Rs.20 brokerage, real-time pricing, and expert portfolio guidance.",
    },
})

# ── 4. Fixed Deposit ─────────────────────────────────────────────────────────
PAGES.append({
    "title": "Fixed Deposit Online India - Best FD Rates 2025 | Shriram",
    "description": "Book a fixed deposit online with up to 8.5% p.a. guaranteed interest. Special senior citizen rates, 80C eligible FD, and flexible tenures. Apply with Shriram today.",
    "hero": {
        "eyebrow": "Fixed Deposits",
        "h1": "Earn Guaranteed Returns with Fixed Deposits",
        "subtitle": "Invest in Fixed Deposits and earn up to 8.5% p.a. guaranteed interest. Safe, predictable, and flexible - with special rates for senior citizens and 80C tax-saving options.",
        "badges": ["Up to 8.5% p.a.", "Capital guaranteed", "Senior citizen extra rate", "80C eligible FD"],
    },
    "sections": [
        {
            "eyebrow": "Types",
            "h2": "Types of Fixed Deposits",
            "body": None,
            "cards": [
                {"title": "Regular Fixed Deposit", "tag": "Most Popular", "desc": "Guaranteed interest on a lump-sum investment for a fixed tenure. Choose cumulative (interest at maturity) or non-cumulative (monthly/quarterly payouts)."},
                {"title": "Tax-Saving FD (ELSS FD)", "tag": "80C Eligible", "desc": "5-year lock-in FD eligible for Rs.1.5 lakh deduction under Section 80C. Interest is taxable; no premature withdrawal allowed during the lock-in."},
                {"title": "Senior Citizen FD", "tag": "Age 60+", "desc": "Extra 0.25%-0.50% interest rate for investors aged 60+. Ideal for generating regular post-retirement income."},
                {"title": "Corporate FD", "tag": "Higher Yields", "desc": "NBFCs like Shriram Finance offer 0.5%-1.5% higher rates than bank FDs. Rated by CRISIL/ICRA - higher yield with assessed credit quality."},
            ],
        },
        {
            "eyebrow": "Interest Rates",
            "h2": "FD Interest Rate Structure",
            "body": "Indicative rates (p.a.). Actual rates may vary. Senior citizens receive an additional 0.50% on all tenures.",
            "table": {
                "headers": ["Tenure", "General Investors", "Senior Citizens"],
                "rows": [
                    ["7 - 45 days", "3.50%", "4.00%"],
                    ["46 - 179 days", "5.50%", "6.00%"],
                    ["180 - 364 days", "6.25%", "6.75%"],
                    ["1 - 2 years", "7.00%", "7.50%"],
                    ["2 - 3 years", "7.15%", "7.65%"],
                    ["3 - 5 years", "7.25%", "7.75%"],
                    ["Above 5 years", "7.00%", "7.50%"],
                ],
            },
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Open a Fixed Deposit",
            "body": None,
            "steps": [
                {"n": "01", "title": "Choose FD Type & Issuer", "desc": "Select between bank FD, NBFC FD, or corporate FD based on your risk tolerance, tenure, and required interest rate."},
                {"n": "02", "title": "Select Tenure & Payout", "desc": "Choose tenure from 7 days to 10 years. Select cumulative (interest at maturity) or non-cumulative (monthly/quarterly payouts)."},
                {"n": "03", "title": "Submit KYC & Apply", "desc": "Submit PAN, Aadhaar, and bank details digitally on the Shriram platform. Application is reviewed and approved in 1-2 days."},
                {"n": "04", "title": "Receive FD Certificate", "desc": "Your FD receipt is issued digitally. Track maturity dates, interest accrual, and renewal options from the Antara app."},
            ],
        },
        {
            "eyebrow": "Comparison",
            "h2": "FD vs Equity vs Mutual Funds",
            "body": None,
            "table": {
                "headers": ["Feature", "Fixed Deposit", "Equity", "Mutual Funds"],
                "rows": [
                    ["Returns", "6.5%-8% p.a. (guaranteed)", "10%-15% (historical, not guaranteed)", "7%-12% (varies by fund type)"],
                    ["Risk", "Very low (guaranteed capital)", "High (market-linked)", "Low to High (depends on fund)"],
                    ["Liquidity", "Low (premature withdrawal penalty)", "High (same-day)", "Medium (T+1 to T+3)"],
                    ["Tax Efficiency", "Low (TDS on interest)", "12.5% LTCG after Rs.1.25L", "Varies by fund type"],
                    ["Best For", "Capital protection, regular income", "Long-term wealth creation", "Goal-based diversified investing"],
                ],
            },
        },
    ],
    "faqs": [
        {"q": "Is FD interest taxable?", "a": "Yes. Interest is taxed per your slab. TDS at 10% is deducted if annual interest exceeds Rs.40,000 (Rs.50,000 for senior citizens). Submit Form 15G or 15H if your income is below the taxable threshold to avoid TDS."},
        {"q": "Can I withdraw my FD before maturity?", "a": "Most FDs allow premature withdrawal with a 0.5%-1% interest penalty. Tax-saving (80C) FDs cannot be broken before the 5-year lock-in. Consider a loan against FD as a penalty-free alternative."},
        {"q": "Is a corporate FD safer than a bank FD?", "a": "Bank FDs are DICGC-insured up to Rs.5 lakh. Corporate FDs aren't covered but offer higher rates. AAA/AA+ rated NBFCs like Shriram Finance have strong track records - assess the rating before investing."},
        {"q": "What is the difference between cumulative and non-cumulative FD?", "a": "Cumulative FDs reinvest interest and pay principal + interest at maturity - ideal for wealth building. Non-cumulative FDs pay interest monthly or quarterly - best for regular income seekers like retirees."},
        {"q": "Can I take a loan against my FD?", "a": "Yes. Borrow up to 70-90% of FD value at 1-2% above the FD rate. Your FD keeps earning interest while you access liquidity - smarter than premature withdrawal."},
    ],
    "cta": {
        "h2": "Secure Your Savings with a Fixed Deposit",
        "body": "Book an FD with Shriram Financial Services - up to 8.5% p.a., fully digital, and with special rates for senior citizens.",
    },
})

# ── 5. NPS ───────────────────────────────────────────────────────────────────
PAGES.append({
    "title": "NPS Investment India - Tax Saving National Pension Scheme | Shriram",
    "description": "Invest in NPS online and save up to Rs.2 lakh in tax deductions annually. Market-linked retirement corpus, PFRDA regulated, flexible fund choice. Open NPS account today.",
    "hero": {
        "eyebrow": "National Pension Scheme",
        "h1": "Retire Rich with NPS - Tax-Efficient & Market-Linked",
        "subtitle": "The National Pension Scheme (NPS) offers up to Rs.2 lakh in tax deductions annually, market-linked returns, and a guaranteed retirement income - all in one PFRDA-regulated account.",
        "badges": ["Up to Rs.2L tax saving", "PFRDA regulated", "Market-linked returns", "Flexible fund choice"],
    },
    "sections": [
        {
            "eyebrow": "Account Types",
            "h2": "NPS Account Types - Tier 1 & Tier 2",
            "body": "NPS operates through two types of accounts designed for different financial needs - mandatory pension savings and flexible voluntary savings.",
            "cards": [
                {"title": "Tier 1 Account", "tag": "Tax-Exempt", "desc": "The primary NPS pension account with tax-deductible contributions. Locked until age 60 - purpose-built for retirement. Minimum: Rs.500/year."},
                {"title": "Tier 2 Account", "tag": "Flexible", "desc": "A voluntary, fully liquid savings account linked to Tier 1. No tax benefit (except for government employees). Minimum: Rs.250/year, withdraw anytime."},
            ],
        },
        {
            "eyebrow": "Asset Classes",
            "h2": "NPS Asset Classes - Where Your Money is Invested",
            "body": None,
            "cards": [
                {"title": "Equity (Class E)", "tag": "Class E", "desc": "Equity and equity-related instruments. Up to 75% allocation (tapers after age 50). Highest growth potential and highest volatility."},
                {"title": "Corporate Bonds (Class C)", "tag": "Class C", "desc": "High-rated corporate debt (AA and above). Moderate risk with stable, predictable returns - good for core stability in your NPS mix."},
                {"title": "Government Securities (Class G)", "tag": "Class G", "desc": "Central and state government bonds - the lowest-risk NPS class. Suitable for conservative investors nearing retirement."},
                {"title": "Alternative Investments (Class A)", "tag": "Class A", "desc": "REITs, InvITs, and alternative assets. Capped at 5% allocation - adds diversification beyond traditional debt and equity."},
            ],
        },
        {
            "eyebrow": "Tax Benefits",
            "h2": "Tax Savings with NPS",
            "body": "NPS offers the highest combined tax deduction among all Section 80 instruments - up to Rs.2 lakh per year with the right combination.",
            "table": {
                "headers": ["Section", "Deduction Limit", "Eligible For", "Note"],
                "rows": [
                    ["80CCD(1)", "Rs.1,50,000 per year", "Employee / Self-employed", "Part of overall 80C limit of Rs.1.5 lakh"],
                    ["80CCD(1B)", "Rs.50,000 per year (additional)", "All NPS subscribers", "Over and above the 80C limit - exclusive to NPS"],
                    ["80CCD(2)", "Up to 10% of Basic + DA", "Salaried (employer contribution)", "No upper limit; employer's contribution is deductible"],
                ],
            },
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Join NPS",
            "body": None,
            "steps": [
                {"n": "01", "title": "Choose PRAN Account Type", "desc": "Select Tier 1 (pension) or Tier 1 + Tier 2 (pension + savings). Choose your Pension Fund Manager (HDFC, SBI, Kotak, etc.)."},
                {"n": "02", "title": "Submit KYC & Documents", "desc": "Complete Aadhaar-based e-KYC on the Shriram platform. Submit PAN, address proof, and nominee details digitally."},
                {"n": "03", "title": "Choose Asset Allocation", "desc": "Select Active Choice (you control the E/C/G split) or Auto Choice (age-based automatic rebalancing) - or the aggressive/moderate/conservative lifecycle fund."},
                {"n": "04", "title": "Start Contributing", "desc": "Make your first contribution (minimum Rs.500 for Tier 1, Rs.250 for Tier 2). Set up auto-debit for regular contributions and track your corpus in real time."},
            ],
        },
        {
            "eyebrow": "Comparison",
            "h2": "NPS vs EPF vs PPF",
            "body": None,
            "table": {
                "headers": ["Feature", "NPS", "EPF", "PPF"],
                "rows": [
                    ["Purpose", "Market-linked retirement savings", "Retirement corpus for salaried", "Tax-free long-term savings"],
                    ["Returns", "Market-linked (8-12% historically)", "8.25% (FY2024, declared by EPFO)", "7.1% (current, reviewed quarterly)"],
                    ["Lock-in", "Until age 60", "Until retirement (partial withdrawal allowed)", "15 years (partial withdrawal from year 7)"],
                    ["Tax on Maturity", "60% tax-free, 40% annuity (taxable)", "Tax-free (if 5+ years)", "Fully tax-free"],
                    ["Employer Contribution", "10% of Basic+DA (govt), 14% for govt employee", "12% of Basic+DA (mandatory)", "No employer contribution"],
                    ["SEBI/PFRDA Regulated", "PFRDA", "EPFO", "Ministry of Finance"],
                ],
            },
        },
    ],
    "faqs": [
        {"q": "Can I withdraw from NPS before retirement?", "a": "Yes - after 3 years, you can withdraw up to 25% of your own contributions for specific reasons: child's education/marriage, home purchase, critical illness, or business setup. A maximum of 3 withdrawals are allowed in a lifetime."},
        {"q": "What happens to NPS at retirement (age 60)?", "a": "At 60, withdraw up to 60% of the corpus tax-free. The remaining 40% must purchase a PFRDA-registered annuity, which pays a monthly pension. You can defer withdrawal until age 75."},
        {"q": "Is NPS better than PPF for tax saving?", "a": "NPS allows up to Rs.2 lakh deduction (80C + 80CCD(1B)), vs PPF's Rs.1.5 lakh under 80C. PPF maturity is fully tax-free; NPS annuity income is taxable. Most advisors recommend both for maximum tax efficiency."},
        {"q": "What is the minimum annual contribution for NPS?", "a": "Tier 1: Rs.500 per contribution, Rs.1,000 per year minimum. Tier 2: Rs.250 per contribution, no annual minimum. Missing Tier 1's annual minimum makes your PRAN inactive with a Rs.100/year penalty."},
        {"q": "Can self-employed individuals invest in NPS?", "a": "Yes. All Indian citizens aged 18-70 can open NPS. Self-employed subscribers claim 80CCD(1) up to 20% of gross income (within Rs.1.5 lakh 80C limit), plus the exclusive Rs.50,000 under 80CCD(1B)."},
    ],
    "cta": {
        "h2": "Secure Your Retirement with NPS Today",
        "body": "Open your NPS PRAN account with Shriram Financial Services. Save tax, grow wealth, and ensure a guaranteed monthly pension.",
    },
})

# ── 6. Global Investing ──────────────────────────────────────────────────────
PAGES.append({
    "title": "Invest US Stocks from India - S&P 500 International ETFs | Shriram",
    "description": "Invest in US stocks, S&P 500, and Nasdaq 100 from India via LRS or India-listed international ETFs. FEMA-compliant, expert-guided global investing. Start today.",
    "hero": {
        "eyebrow": "Global Investing",
        "h1": "Invest in US Stocks & Global Markets from India",
        "subtitle": "Diversify beyond India. Access S&P 500, Nasdaq 100, and global leaders like Apple, NVIDIA, and Microsoft - with expert guidance and LRS-compliant processes from Shriram Financial Services.",
        "badges": ["USD 250,000 LRS limit", "No LRS for India-listed ETFs", "Expert global research", "Currency & FX guidance"],
    },
    "sections": [
        {
            "eyebrow": "Global Markets",
            "h2": "Markets You Can Access",
            "body": None,
            "cards": [
                {"title": "United States", "tag": "S&P 500, Nasdaq 100, Dow Jones", "desc": "The world's largest and most liquid market. Access Apple, Microsoft, Amazon, NVIDIA, and thousands of global category leaders with deep research coverage."},
                {"title": "Europe", "tag": "FTSE 100, DAX, CAC 40, STOXX 50", "desc": "Established multinationals in luxury, automotive, pharma, and banking. Built-in currency diversification across EUR, GBP, and CHF."},
                {"title": "Asia Pacific", "tag": "Hang Seng, Nikkei 225, MSCI Asia", "desc": "Chinese tech, Japanese precision manufacturing, and South Korean semiconductors - high-growth markets with distinct economic cycles from India."},
            ],
        },
        {
            "eyebrow": "How It Works",
            "h2": "Ways to Invest Internationally from India",
            "body": None,
            "cards": [
                {"title": "LRS (Liberalised Remittance Scheme)", "desc": "Remit up to USD 250,000 per financial year for overseas investment under FEMA's LRS. Funds go to your foreign brokerage account or via India-based fund-of-funds."},
                {"title": "India-Listed Global Funds", "desc": "The simplest route - buy Nasdaq 100 ETFs or S&P 500 index funds listed on Indian exchanges. No LRS needed; invest like any domestic mutual fund."},
                {"title": "Direct Foreign Stock Platform", "desc": "Buy fractional US shares (Apple, NVIDIA, Berkshire) directly via LRS-linked platforms. More flexibility for individual stock-pickers."},
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Start Global Investing",
            "body": None,
            "steps": [
                {"n": "01", "title": "Choose Your Route", "desc": "Decide between India-listed international ETFs/funds (no LRS, simplest) or a direct foreign stock account (LRS required, more flexibility)."},
                {"n": "02", "title": "Complete KYC & Currency Exchange", "desc": "For LRS-based investing, submit Form A2 to your bank and complete FEMA compliance. For India-listed funds, standard KYC through Shriram is sufficient."},
                {"n": "03", "title": "Select Markets & Instruments", "desc": "Choose geographic markets, index funds, ETFs, or specific stocks. Our research desk provides guidance on US and international market opportunities."},
                {"n": "04", "title": "Monitor Global Portfolio", "desc": "Track your international positions alongside domestic holdings on the Antara app. Set FX rate alerts and portfolio rebalancing reminders."},
            ],
        },
        {
            "eyebrow": "Know Before You Invest",
            "h2": "Risks of Global Investing",
            "body": None,
            "risks": [
                {"label": "Currency Risk", "desc": "Returns in USD or EUR convert to INR on exit. If INR strengthens, your rupee returns fall even if the underlying stock rose."},
                {"label": "Regulatory Risk", "desc": "RBI and SEBI periodically revise LRS limits and fund-of-fund structures. A regulatory change can restrict overseas investment or affect existing holdings."},
                {"label": "Market & Geopolitical Risk", "desc": "Trade wars, sanctions, and foreign central bank policy drive international markets - factors outside India-based investors' control."},
                {"label": "TCS on Remittance", "desc": "LRS remittances above Rs.7 lakh/year attract 20% TCS under Section 206C(1G). It's refundable via ITR, but creates temporary cash flow drag."},
            ],
        },
    ],
    "faqs": [
        {"q": "Can Indian residents invest in US stocks?", "a": "Yes - two routes: (1) India-listed international ETFs/funds, no LRS needed; (2) Direct US stock investment via a foreign brokerage linked to an LRS-compliant Indian bank, allowing fractional shares."},
        {"q": "What is the LRS limit for investing abroad?", "a": "Under FEMA's LRS, Indian residents can remit up to USD 250,000 (~Rs.2.08 crore) per financial year for overseas investment. Remittances above Rs.7 lakh attract 20% TCS, refundable via ITR."},
        {"q": "How are international investment returns taxed in India?", "a": "Direct LRS holdings: short-term gains (under 2 years) taxed per slab; long-term gains (2+ years) at 12.5%. India-listed international ETFs are taxed as debt funds - per slab for both short and long term (from April 2023)."},
        {"q": "Is investing in international markets risky for Indian investors?", "a": "Global diversification reduces overall portfolio correlation - US and Indian markets don't always move together. Currency, geopolitical, and regulatory risks are additional. A 10-20% allocation via India-listed international ETFs suits most retail investors."},
    ],
    "cta": {
        "h2": "Go Global with Shriram Financial Services",
        "body": "Access US stocks, global ETFs, and international funds - expert-guided, FEMA-compliant, and fully digital.",
    },
})

# ── 7. Insurance ─────────────────────────────────────────────────────────────
PAGES.append({
    "title": "Term Health Insurance Online India - Compare Plans | Shriram",
    "description": "Compare term life, health, motor, and ULIP insurance plans from India's top IRDAI-regulated insurers. Expert advisory, fast digital application. Get covered today.",
    "hero": {
        "eyebrow": "Insurance",
        "h1": "Protect What Matters Most - Life, Health & More",
        "subtitle": "Compare term life, health, motor, and ULIP insurance plans from India's leading insurers. Expert advisory, easy digital application, and dedicated claims support from Shriram Financial Services.",
        "badges": ["Multiple insurer comparison", "IRDAI-regulated partners", "Tax benefits (80C, 80D)", "Claims assistance"],
    },
    "sections": [
        {
            "eyebrow": "Products",
            "h2": "Insurance Products We Offer",
            "body": "We partner with India's leading IRDAI-regulated insurance companies to offer comprehensive coverage for every stage of life.",
            "cards": [
                {"title": "Term Life Insurance", "highlight": "Highest cover at lowest premium", "desc": "Pure life cover for a fixed term - nominee receives sum assured on death. No maturity benefit, but the highest cover at the lowest premium."},
                {"title": "Health Insurance", "highlight": "Cashless treatment at 10,000+ hospitals", "desc": "Covers hospitalisation, surgery, ICU, daycare, and pre/post-hospitalisation costs. Shields your savings from catastrophic medical bills."},
                {"title": "Motor Insurance", "highlight": "Mandatory by law (third-party)", "desc": "Mandatory third-party cover for all vehicles in India. Comprehensive plans add own-damage, theft, and natural calamity protection."},
                {"title": "ULIP", "highlight": "Insurance + investment in one", "desc": "Unit Linked Insurance Plan - part of the premium provides life cover, and the remainder invests in market-linked equity/debt funds."},
            ],
        },
        {
            "eyebrow": "Why Insurance",
            "h2": "Why You Need Insurance",
            "body": None,
            "cards": [
                {"title": "Protect Your Family", "desc": "Life cover ensures your family can maintain their lifestyle, repay home loans, and fund children's education - even in your absence."},
                {"title": "Shield Against Medical Costs", "desc": "One hospitalisation can cost Rs.2-10 lakh. Health insurance prevents a medical emergency from wiping out your savings and investments."},
                {"title": "Save on Tax", "desc": "Life insurance premiums get Rs.1.5 lakh deduction under 80C. Health premiums get up to Rs.50,000 under 80D for senior citizens."},
                {"title": "Mandatory Compliance", "desc": "Motor third-party insurance is legally required for every vehicle. Lapsed cover means fines and personal liability in accidents."},
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Choose & Buy Insurance",
            "body": None,
            "steps": [
                {"n": "01", "title": "Assess Your Needs", "desc": "Determine the type and quantum of cover needed. Use our online calculators: HLV (Human Life Value) for life insurance, sum insured calculator for health."},
                {"n": "02", "title": "Compare Plans", "desc": "Compare premiums, features, claim settlement ratios, and network hospitals across multiple insurance providers on the Shriram platform."},
                {"n": "03", "title": "Apply & Undergo Medical Tests", "desc": "Fill the proposal form digitally. For higher life insurance covers, an insurer-appointed doctor may conduct a medical examination."},
                {"n": "04", "title": "Receive Policy Document", "desc": "Once underwriting is complete and premium is paid, your policy document is issued digitally. Claim support is provided by Shriram's dedicated insurance team."},
            ],
        },
    ],
    "faqs": [
        {"q": "How much life insurance cover do I need?", "a": "A standard benchmark is 10-15x your annual income. Factor in home loan balances, children's education costs, and spouse's income gap. Use Shriram's HLV calculator for a precise figure."},
        {"q": "What is the difference between a cashless and reimbursement health claim?", "a": "Cashless claims are settled directly between your insurer and a network hospital - you pay nothing upfront. Reimbursement claims require you to pay the bill first and apply for refund; used for non-network hospitals."},
        {"q": "Does health insurance cover pre-existing diseases?", "a": "Most plans have a 2-4 year waiting period for pre-existing diseases. After the waiting period, they're covered. Always disclose PEDs accurately - non-disclosure leads to claim rejection."},
        {"q": "Is it better to buy a term plan or an endowment plan?", "a": "Term plans deliver more cover per rupee. Endowment and ULIP returns often trail a term plan + mutual fund combination. Most advisors recommend 'buy term, invest the rest' for optimal outcomes."},
        {"q": "What is the claim settlement ratio (CSR) and why does it matter?", "a": "CSR is the % of claims settled by an insurer in a year. IRDAI publishes CSR annually. Above 97% signals reliability - Shriram recommends only partners with strong CSR track records."},
    ],
    "cta": {
        "h2": "Protect Your Family's Future Today",
        "body": "Compare India's best insurance plans and get expert guidance from Shriram Financial Services. Fast application, transparent pricing, dedicated claims support.",
    },
})

# ── 8. Commodities ───────────────────────────────────────────────────────────
PAGES.append({
    "title": "Commodity Trading India - MCX Gold Silver Crude Oil | Shriram",
    "description": "Trade gold, silver, crude oil, and agricultural commodities on MCX and NCDEX. Flat Rs.20 brokerage, research-backed advisory, and physical delivery support. Start today.",
    "hero": {
        "eyebrow": "Commodities",
        "h1": "Trade Gold, Crude Oil & Agricultural Commodities",
        "subtitle": "Diversify your portfolio with commodity futures and options on MCX and NCDEX. Hedge against inflation, trade energy cycles, and capitalise on agri seasonality.",
        "badges": ["MCX & NCDEX Access", "Flat Rs.20/order", "Research-backed", "Physical delivery support"],
    },
    "sections": [
        {
            "eyebrow": "Asset Classes",
            "h2": "Types of Commodities You Can Trade",
            "body": "Commodities are broadly split into four categories. Each behaves differently and is driven by distinct supply-demand dynamics.",
            "cards": [
                {"title": "Precious Metals", "tag": "Gold (GOLD), Silver (SILVER), Platinum", "desc": "Safe-haven assets that appreciate during inflation, currency weakness, or market uncertainty."},
                {"title": "Energy", "tag": "Crude Oil (CRUDEOIL), Natural Gas (NATURALGAS)", "desc": "Driven by global supply-demand, OPEC policy, and geopolitical events. Highly liquid with active participation."},
                {"title": "Agricultural", "tag": "Cotton, Chana, Castor Seed, Jeera, Guar", "desc": "Seasonally influenced by monsoon, production cycles, and export demand. Traded on NCDEX with physical delivery."},
                {"title": "Base Metals", "tag": "Copper, Aluminium, Zinc, Nickel, Lead", "desc": "Tied to industrial output and infrastructure demand. Widely used as leading indicators of global economic growth."},
            ],
        },
        {
            "eyebrow": "Exchanges",
            "h2": "How the Commodity Market Works in India",
            "body": None,
            "cards": [
                {"title": "MCX - Multi Commodity Exchange", "desc": "India's largest commodity exchange with 90%+ market share. Focuses on metals and energy. Also offers options on Gold, Silver, and Crude Oil."},
                {"title": "NCDEX - National Commodity & Derivatives Exchange", "desc": "India's premier agri commodity exchange. Futures and options on chana, castor seed, guar, and cotton - with physical delivery settlement."},
            ],
        },
        {
            "eyebrow": "Requirements",
            "h2": "Eligibility & Documents Required",
            "body": None,
            "eligibility": [
                "Indian resident individual, HUF, or corporate entity",
                "Valid PAN Card (mandatory)",
                "Active trading account with commodity segment enabled",
                "Bank account for fund settlement",
                "For physical delivery: warehouse receipt and delivery centre access",
            ],
            "documents": [
                "PAN Card (mandatory)",
                "Address Proof - Aadhaar / Passport / Utility Bill",
                "Bank Account Details - cancelled cheque or bank statement",
                "Photograph (passport-size, JPEG format)",
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Start Trading Commodities",
            "body": None,
            "steps": [
                {"n": "01", "title": "Enable Commodity Segment", "desc": "Add MCX/NCDEX commodity trading segment to your existing Shriram trading account - no new account needed."},
                {"n": "02", "title": "Add Margin Funds", "desc": "Transfer margin funds for the contracts you wish to trade. Commodity futures require SPAN + Exposure margin."},
                {"n": "03", "title": "Choose Commodity & Contract", "desc": "Select the commodity, lot size, and expiry month. Check live commodity prices and research before entering."},
                {"n": "04", "title": "Trade & Monitor", "desc": "Place buy or sell orders through the Antara platform. Monitor live spot prices, basis, and roll costs."},
            ],
        },
        {
            "eyebrow": "Costs",
            "h2": "Brokerage & Charges",
            "body": None,
            "table": {
                "headers": ["Charge Type", "Amount / Rate"],
                "rows": [
                    ["Brokerage", "Rs.20 per executed order"],
                    ["CTT (Commodity Transaction Tax)", "0.01% on sell side (non-agri); Nil on agri futures"],
                    ["GST", "18% on brokerage + exchange charges"],
                    ["Exchange Charges (MCX)", "0.0026% (metals); 0.00035% (energy)"],
                    ["SEBI Turnover Fee", "Rs.10 per crore"],
                ],
            },
        },
        {
            "eyebrow": "Know Before You Trade",
            "h2": "Risks of Commodity Trading",
            "body": None,
            "risks": [
                {"label": "Price Volatility", "desc": "Commodity prices are influenced by global macroeconomic factors, weather events, and geopolitical tensions - causing sudden, large price swings."},
                {"label": "Delivery Risk", "desc": "If a futures position is not squared off before expiry, it may result in compulsory physical delivery - requiring warehouse receipts and logistics."},
                {"label": "Basis Risk", "desc": "The difference between spot price and futures price (basis) can widen unexpectedly, reducing hedge effectiveness and increasing unexpected mark-to-market losses."},
                {"label": "Leverage Risk", "desc": "Commodity futures are traded on margin. A small adverse price movement can result in losses much larger than the initial margin deposited."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is the minimum investment to start commodity trading in India?", "a": "No regulatory minimum, but you need margin to cover one contract. GOLD Mini (100g) requires Rs.20,000-Rs.40,000; Crude Oil (100 barrels) requires Rs.15,000-Rs.30,000 depending on volatility."},
        {"q": "What are the trading hours for commodities on MCX?", "a": "MCX trades Mon-Fri: day session 9:00 AM-5:00 PM; evening session 5:00 PM-11:30 PM (11:55 PM during US daylight saving). NCDEX agri commodities trade 9:00 AM-5:00 PM only."},
        {"q": "Is physical delivery mandatory for commodity futures?", "a": "Only if you hold through the tender period (last few days before expiry). Most retail traders exit before this window. Shriram's platform alerts you as expiry nears so you can roll or close."},
        {"q": "How are commodity trades settled?", "a": "Futures are marked-to-market daily - P&L is credited/debited each day. At expiry, settlement is either physical delivery or cash settlement at the daily settlement price."},
        {"q": "Can I hedge my business exposure using commodity futures?", "a": "Yes. MCX and NCDEX futures let jewellers, oil distributors, and agri traders lock in input prices. Shriram's research team provides hedging advisory to eligible clients."},
    ],
    "cta": {
        "h2": "Start Commodity Trading Today",
        "body": "Access MCX and NCDEX with a single account. Flat Rs.20 brokerage, live prices, and expert research support.",
    },
})

# ── 9. Currency ──────────────────────────────────────────────────────────────
PAGES.append({
    "title": "Currency Trading India - USD INR Forex Futures NSE | Shriram",
    "description": "Trade USD/INR, EUR/INR, GBP/INR, and JPY/INR currency futures on NSE. Flat Rs.20 brokerage, low margin, expert guidance. Enable forex trading with Shriram today.",
    "hero": {
        "eyebrow": "Currency",
        "h1": "Trade Currency Futures & Hedge Forex Risk",
        "subtitle": "Access all four INR currency pairs - USD, EUR, GBP, and JPY - with live pricing, low margin requirements, and expert guidance from Shriram Financial Services.",
        "badges": [],
    },
    "sections": [
        {
            "eyebrow": "Currency Pairs",
            "h2": "Currency Pairs Available on NSE",
            "body": "India's NSE offers futures and options on four INR-based currency pairs. All are cash-settled in INR at the RBI reference rate on expiry.",
            "currency_pairs": [
                {"pair": "USD/INR", "full": "US Dollar / Indian Rupee", "desc": "India's most liquid currency pair. Moves on RBI interventions, US Fed policy, and India's trade balance data.", "lot": "1,000 USD", "tick": "0.25 paise"},
                {"pair": "EUR/INR", "full": "Euro / Indian Rupee", "desc": "Driven by ECB policy and Eurozone data. Second most traded INR pair on NSE.", "lot": "1,000 EUR", "tick": "0.25 paise"},
                {"pair": "GBP/INR", "full": "British Pound / Indian Rupee", "desc": "Reacts to Bank of England decisions and UK economic indicators.", "lot": "1,000 GBP", "tick": "0.25 paise"},
                {"pair": "JPY/INR", "full": "Japanese Yen / Indian Rupee", "desc": "Driven by Bank of Japan policy and risk-on/risk-off flows. Priced per 100 JPY.", "lot": "1,00,000 JPY", "tick": "0.25 paise"},
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Start Currency Trading",
            "body": None,
            "steps": [
                {"n": "01", "title": "Enable Currency Segment", "desc": "Add the currency derivatives segment to your Shriram trading account - fully digital, no branch visit required."},
                {"n": "02", "title": "Add Trading Margin", "desc": "Transfer the required SPAN + Exposure margin. Currency futures require relatively low margins compared to equity F&O."},
                {"n": "03", "title": "Choose Pair & Expiry", "desc": "Select your currency pair (USD/INR, EUR/INR, etc.) and the contract expiry - monthly contracts settling on the last business day."},
                {"n": "04", "title": "Trade & Hedge", "desc": "Place orders, monitor your P&L in real time, and use the position to either speculate or hedge your forex exposure."},
            ],
        },
        {
            "eyebrow": "Costs",
            "h2": "Brokerage & Charges",
            "body": None,
            "table": {
                "headers": ["Charge Type", "Amount / Rate"],
                "rows": [
                    ["Brokerage", "Rs.20 per executed order"],
                    ["STT", "Nil (not applicable on currency derivatives)"],
                    ["GST", "18% on brokerage + exchange charges"],
                    ["Exchange Charges (NSE Currency)", "0.00035%"],
                    ["SEBI Turnover Fee", "Rs.1 per crore"],
                ],
            },
        },
        {
            "eyebrow": "Comparison",
            "h2": "Currency Derivatives vs Equity Derivatives",
            "body": None,
            "table": {
                "headers": ["Feature", "Currency Derivatives", "Equity Derivatives"],
                "rows": [
                    ["Purpose", "Hedge forex risk or speculate on exchange rates", "Invest in company ownership"],
                    ["Underlying", "Currency pair (USD/INR, EUR/INR, etc.)", "Stocks, indices"],
                    ["Lot Size", "1,000 units of foreign currency", "Varies per contract"],
                    ["Settlement", "Cash settlement in INR at RBI reference rate", "Physical or cash settlement"],
                    ["Margin Required", "Low (2-4% of contract value)", "Higher (5-15%)"],
                    ["Trading Hours", "9:00 AM - 5:00 PM (Mon-Fri)", "9:15 AM - 3:30 PM"],
                ],
            },
        },
        {
            "eyebrow": "Know Before You Trade",
            "h2": "Risks of Currency Trading",
            "body": None,
            "risks": [
                {"label": "Exchange Rate Risk", "desc": "Currency prices are highly sensitive to central bank policy, macroeconomic data releases, and geopolitical events - causing sudden, large moves."},
                {"label": "RBI Intervention Risk", "desc": "The Reserve Bank of India actively intervenes in the INR market to curb excess volatility - creating sharp, unpredictable counter-trend moves."},
                {"label": "Leverage Risk", "desc": "Currency futures are leveraged instruments. A small adverse rate movement can result in margin calls and losses exceeding your initial deposit."},
                {"label": "Liquidity Risk", "desc": "EUR/INR, GBP/INR, and JPY/INR pairs have significantly lower liquidity than USD/INR - making it harder to exit large positions at fair prices."},
            ],
        },
    ],
    "faqs": [
        {"q": "Who should trade currency derivatives in India?", "a": "Exporters and importers hedge forex receivables/payables. Portfolio investors hedge foreign asset exposure. Retail traders with macro knowledge can speculate on exchange rate movements."},
        {"q": "Is RBI approval required for currency trading?", "a": "No special approval is needed for residents trading on NSE/BSE. For speculative USD/INR positions up to $100 million, no underlying forex exposure is required."},
        {"q": "How is currency futures settlement done?", "a": "Currency futures are cash-settled in INR. On expiry (last business day of the month), settlement uses the RBI reference rate published at 12:30 PM. No physical forex delivery occurs."},
        {"q": "What is the USDINR lot size?", "a": "One USD/INR lot = 1,000 USD, with a notional value of ~Rs.83,000-Rs.85,000 at current rates. Non-bank entities can hold up to 6 million USD in open positions."},
        {"q": "Can NRIs trade currency derivatives in India?", "a": "No. Only persons resident in India as defined by FEMA are eligible. NRIs are currently not permitted to participate in Indian currency futures and options markets."},
    ],
    "cta": {
        "h2": "Trade Currency Derivatives Today",
        "body": "Hedge your forex exposure or speculate on exchange rates. Flat Rs.20 brokerage, live pricing, and expert support.",
    },
})

# ── 10. Loans ────────────────────────────────────────────────────────────────
PAGES.append({
    "title": "Personal Loan Online India - Fast Approval Loan Against Securities | Shriram",
    "description": "Apply for personal loans, business loans, and loans against mutual funds or stocks online. Approval in 24 hours, competitive rates, minimal paperwork. Apply with Shriram.",
    "hero": {
        "eyebrow": "Loans",
        "h1": "Access Instant Funds Without Selling Your Investments",
        "subtitle": "From personal loans to loans against your portfolio - get liquidity when you need it, at competitive rates, with minimal paperwork from Shriram Financial Services.",
        "badges": ["Approval in 24-48 hours", "No collateral (personal)", "Portfolio stays invested", "Competitive rates"],
    },
    "sections": [
        {
            "eyebrow": "Our Loan Products",
            "h2": "Loan Solutions for Every Need",
            "body": "Whether you need liquidity for personal use, business growth, or want to unlock value from your investment portfolio - Shriram has a loan product for you.",
            "loan_cards": [
                {"title": "Personal Loan", "rate": "11%-24% p.a.", "amount": "Up to Rs.40 Lakhs", "desc": "Unsecured loan for medical emergencies, travel, home renovation, or debt consolidation. Fast digital approval with minimal paperwork."},
                {"title": "Loan Against Mutual Funds", "rate": "9.5%-14% p.a.", "amount": "Up to Rs.1 Crore", "desc": "Pledge mutual fund units as collateral - access up to 50-80% of NAV as a loan. SIPs keep running and units stay invested."},
                {"title": "Loan Against Stocks", "rate": "10.5%-15% p.a.", "amount": "Up to Rs.5 Crore", "desc": "Pledge your equity portfolio for instant liquidity - without selling a single share. Dividends and appreciation continue while shares are pledged."},
                {"title": "Business Loan", "rate": "13%-22% p.a.", "amount": "Up to Rs.50 Lakhs", "desc": "Unsecured MSME loans for working capital, equipment, or expansion. Fast processing for self-employed professionals and growing businesses."},
            ],
        },
        {
            "eyebrow": "Requirements",
            "h2": "Eligibility & Documents Required",
            "body": None,
            "eligibility": [
                "Indian resident aged 21-65 years (personal) or 21-70 years (business)",
                "Minimum monthly income: Rs.20,000 (salaried) or Rs.25,000 (self-employed)",
                "Credit score (CIBIL) of 700+ recommended for best rates",
                "Stable employment of 6+ months (salaried) or 2+ years in business",
                "Active bank account with regular transaction history",
            ],
            "documents": [
                "PAN Card and Aadhaar Card",
                "Address Proof - Aadhaar / Passport / Utility Bill",
                "Income Proof - last 3 months salary slips or ITR (2 years)",
                "Bank statements (last 6 months)",
                "Business proof (for business loans) - GST certificate, trade licence",
            ],
        },
        {
            "eyebrow": "Getting Started",
            "h2": "How to Apply for a Loan",
            "body": None,
            "steps": [
                {"n": "01", "title": "Check Eligibility", "desc": "Use the online eligibility calculator. Enter your income, existing EMIs, and loan amount needed to see your options instantly."},
                {"n": "02", "title": "Submit Application", "desc": "Fill the digital application form on the Shriram platform. Upload documents via the app - no branch visit needed."},
                {"n": "03", "title": "Credit Assessment", "desc": "Our team verifies income, credit history, and repayment capacity. Approval decision is typically communicated within 24-48 hours."},
                {"n": "04", "title": "Disbursal", "desc": "Loan amount credited directly to your bank account. For secured loans (against MF/stocks), pledge is registered digitally before disbursal."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is the difference between Loan Against MF and Loan Against Stocks?", "a": "LAMF lets you pledge mutual fund units for up to 50-80% of NAV. LAS lets you pledge listed equity shares for up to 50% of market value. Both keep your investments earning returns while you access liquidity."},
        {"q": "Does a personal loan from Shriram require a salary slip?", "a": "Salaried applicants need the last 3 months' salary slips and bank statements. Self-employed applicants can use 2 years' ITR and 6-month bank statements. Informal income programs are also available - contact your Shriram branch."},
        {"q": "What happens if I miss an EMI?", "a": "Missing an EMI triggers a late charge and a negative CIBIL entry. For LAMF/LAS, if the LTV breaches the threshold, Shriram may issue a margin call. Persistent default can result in pledged securities being sold."},
        {"q": "Can I prepay my loan before the tenure ends?", "a": "Yes. Personal and business loans can be prepaid after 6-12 months lock-in, with 2-5% foreclosure charges. Secured loans against MF/stocks can be repaid free after 3 months."},
        {"q": "How is interest calculated on a Loan Against Securities?", "a": "LAMF and LAS use overdraft-style interest - you pay only on the amount drawn, not the total approved limit. This makes it far cheaper than a term loan when you need partial access."},
    ],
    "cta": {
        "h2": "Get the Funds You Need - Fast",
        "body": "Apply for a personal loan, business loan, or unlock value from your portfolio - all fully digital with 24-hour approval.",
    },
})

# ── 11. LAMF ─────────────────────────────────────────────────────────────────
PAGES.append({
    "title": "Loan Against Mutual Funds India - LAMF Pledge MF | Shriram",
    "description": "Get instant cash by pledging mutual fund units as collateral. No redemption required, SIPs keep running, low interest rates from 9.5% p.a. Apply for LAMF today.",
    "hero": {
        "eyebrow": "Loan Against Mutual Funds",
        "h1": "Unlock Cash from Your Mutual Fund Portfolio",
        "subtitle": "Need funds urgently? Pledge your mutual fund units and get instant liquidity - without selling your investments or stopping your SIPs.",
        "badges": ["No Redemption", "Low Interest Rates", "SIPs Continue", "Overdraft Facility"],
    },
    "sections": [
        {
            "eyebrow": None,
            "h2": "Why Choose LAMF?",
            "body": None,
            "cards": [
                {"title": "No Redemption Needed", "desc": "Pledge your mutual fund units as collateral without selling them. NAV growth continues on all pledged holdings."},
                {"title": "Instant Liquidity", "desc": "Funds credited to your bank within 24-48 hours of pledge creation - no branch visits, no paper queues."},
                {"title": "Low Interest Rates", "desc": "LAMF rates start at 9.5% p.a. - far cheaper than personal loans or credit cards because your funds secure the loan."},
                {"title": "Overdraft Facility", "desc": "Draw what you need and pay interest only on the amount withdrawn. Repay and redraw anytime within your limit."},
                {"title": "SIPs Continue Running", "desc": "Your SIP deductions and new unit additions continue after pledging. Long-term compounding stays on track."},
                {"title": "No Prepayment Charges", "desc": "Repay partially or fully whenever you have surplus - no foreclosure penalty after the initial lock-in period."},
            ],
        },
        {
            "eyebrow": None,
            "h2": "Eligible Fund Types & LTV Ratios",
            "body": "The loan amount you receive depends on the Loan-to-Value ratio of the mutual funds you pledge.",
            "table": {
                "headers": ["Fund Type", "LTV Ratio", "Examples"],
                "rows": [
                    ["Equity Mutual Funds", "50-60%", "Large-cap, mid-cap, flexi-cap, ELSS funds"],
                    ["Debt Mutual Funds", "70-80%", "Liquid funds, short-term bond funds, gilt funds"],
                    ["Hybrid / Balanced Funds", "55-65%", "Balanced advantage, aggressive hybrid funds"],
                    ["Index Funds / ETFs", "50-60%", "NIFTY 50 index funds, SENSEX ETFs"],
                ],
            },
        },
        {
            "eyebrow": None,
            "h2": "How to Get a Loan Against Mutual Funds",
            "body": None,
            "steps": [
                {"n": "01", "title": "Link Your Mutual Fund Holdings", "desc": "Connect your existing mutual fund portfolio to your Shriram account. Folios held across AMCs are visible in one dashboard."},
                {"n": "02", "title": "Select Units to Pledge", "desc": "Choose the fund(s) you want to pledge. The system calculates your eligible loan amount (LTV) based on current NAV."},
                {"n": "03", "title": "Submit Pledge Request", "desc": "Submit the pledge digitally - the AMC registers the lien on your units with the RTA. No physical forms required."},
                {"n": "04", "title": "Receive Loan in Your Account", "desc": "Once the pledge is confirmed, funds are credited to your linked bank account. Draw as little or as much as you need."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is Loan Against Mutual Funds (LAMF)?", "a": "LAMF is a secured loan where you pledge MF units as collateral. The lender places a lien via the RTA - you retain ownership and units keep earning returns, but cannot be redeemed until the lien is released on repayment."},
        {"q": "How much loan can I get against my mutual funds?", "a": "LTV varies by fund type: equity funds 50-60%, debt funds 70-80%, hybrid funds 55-65%. Pledge equity funds worth Rs.10 lakh and receive Rs.5-6 lakh in credit."},
        {"q": "Do my SIPs continue during the loan period?", "a": "Yes. Only the specific units pledged are locked. New SIP units added each month are not auto-pledged - your SIP corpus keeps growing alongside the pledged holdings."},
        {"q": "What happens if the NAV falls below the loan value?", "a": "If NAV drops and the loan exceeds the permissible LTV, Shriram issues a margin call asking for more collateral or partial repayment. Real-time LTV monitoring alerts you before forced action is needed."},
        {"q": "Can I repay the LAMF loan early?", "a": "Yes - LAMF is an overdraft facility. Repay any amount at any time without penalty after the initial lock-in (typically 3 months). Full repayment releases the pledge lien immediately."},
    ],
    "cta": {
        "h2": "Ready to Unlock Your Portfolio's Value?",
        "body": "Get instant liquidity at low interest rates without disrupting your investments.",
    },
})

# ── 12. LAS ──────────────────────────────────────────────────────────────────
PAGES.append({
    "title": "Loan Against Shares India - Pledge Stocks for Loan LAS | Shriram",
    "description": "Pledge your equity shares for instant liquidity without selling. Low rates from 10.5% p.a., flexible overdraft, dividends continue. Apply for LAS with Shriram today.",
    "hero": {
        "eyebrow": "Loan Against Stocks",
        "h1": "Borrow Against Your Share Portfolio Instantly",
        "subtitle": "Pledge your equity holdings and unlock a flexible credit line - without selling a single share. Your stocks stay, your dividends continue, and you get the funds you need.",
        "badges": ["No Share Sale", "Keep Dividends", "Flexible Overdraft", "Low Interest Rates"],
    },
    "sections": [
        {
            "eyebrow": None,
            "h2": "Why Choose LAS?",
            "body": None,
            "cards": [
                {"title": "No Need to Sell Shares", "desc": "Pledge stocks as collateral without selling - shares stay in your demat account, earning dividends and bonus shares throughout the loan."},
                {"title": "High Loan Amount", "desc": "Borrow up to 50-60% of pledged share value. Large portfolios unlock significant liquidity without any asset liquidation."},
                {"title": "Flexible Overdraft", "desc": "Draw funds when needed, repay when convenient, and redraw again - pay interest only on the outstanding amount, not the total limit."},
                {"title": "Low Interest Rates", "desc": "LAS rates start at 10% p.a. - far below unsecured personal loans and credit cards, because your shares secure the lender."},
                {"title": "Continue Earning Benefits", "desc": "Dividends, bonus shares, and rights issues on pledged stocks are all credited to your demat account during the loan period."},
                {"title": "Digital Pledge Process", "desc": "Pledge creation is 100% online via CDSL/NSDL. No physical certificates or branch visits - done within minutes."},
            ],
        },
        {
            "eyebrow": None,
            "h2": "Eligible Securities & LTV Ratios",
            "body": "Loan limits depend on the type and market value of securities pledged. Only exchange-listed, lender-approved shares qualify.",
            "table": {
                "headers": ["Security Type", "LTV Ratio", "Examples"],
                "rows": [
                    ["BSE/NSE Listed Shares", "50-60%", "NIFTY 50, NIFTY 100, large-cap stocks"],
                    ["Mid-cap Stocks", "40-50%", "NIFTY Midcap 150 constituents (lender-approved list)"],
                    ["ETFs (Equity)", "50-60%", "Nifty BeES, SBI ETF Nifty 50"],
                    ["Sovereign Gold Bonds", "60-70%", "RBI-issued SGBs held in demat form"],
                ],
            },
        },
        {
            "eyebrow": None,
            "h2": "How to Get a Loan Against Stocks",
            "body": None,
            "steps": [
                {"n": "01", "title": "Submit Pledge Request", "desc": "Select the stocks you want to pledge from your demat account. Shriram initiates a pledge request to CDSL/NSDL on your behalf."},
                {"n": "02", "title": "Authorise via CDSL/NSDL", "desc": "You receive a TPIN-based authorisation request on your registered mobile and email. One-time approval authorises the pledge."},
                {"n": "03", "title": "Pledge Confirmed & Limit Set", "desc": "Once CDSL/NSDL confirms the pledge mark, your loan limit is activated based on the LTV of pledged shares."},
                {"n": "04", "title": "Withdraw & Repay Anytime", "desc": "Draw funds to your bank account. Repay whenever convenient. The pledge is released when the loan is fully repaid."},
            ],
        },
    ],
    "faqs": [
        {"q": "What is Loan Against Stocks (LAS)?", "a": "LAS is a secured overdraft where you pledge equity shares or ETFs in your demat account. CDSL/NSDL marks a lien on the pledged securities - you retain ownership and dividends, but cannot sell until the lien is lifted on full repayment."},
        {"q": "Which stocks are eligible for pledging?", "a": "Lenders accept NIFTY 50/100 constituents, large-cap and mid-cap exchange-listed shares, and ETFs from their approved list. Penny stocks, illiquid shares, and stocks under special surveillance are ineligible."},
        {"q": "Do I still receive dividends on pledged shares?", "a": "Yes - pledging transfers only a lien, not ownership. Dividends, bonus shares, rights, and stock splits are credited to your demat account throughout the loan period."},
        {"q": "What happens if the share price falls sharply?", "a": "If pledged value drops and the loan exceeds the LTV (typically 50-60%), Shriram issues a margin call. You must pledge more shares, deposit cash, or repay part of the loan. The system alerts you well before forced action is triggered."},
        {"q": "How is LAS different from Margin Trading Facility (MTF)?", "a": "MTF is specifically for buying more securities on margin in your trading account - the bought shares serve as collateral. LAS is a general-purpose credit line against existing demat holdings, disbursed to your bank account for any use, typically at lower rates and longer tenors."},
    ],
    "cta": {
        "h2": "Make Your Portfolio Work Harder",
        "body": "Access funds without selling your shares. Keep your portfolio intact and growing.",
    },
})


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Title Page ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Shriram Financial Services")
    run.bold = True
    run.font.size = Pt(28)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_para.add_run("Product Pages Content")
    run2.font.size = Pt(18)

    doc.add_paragraph()
    desc_para = doc.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_para.add_run("Complete text content extracted from all 12 product pages.")

    # ── Process each page ─────────────────────────────────────────────────────
    for page in PAGES:
        add_page_break(doc)

        # Page title as H1
        h1(doc, page["title"])
        body(doc, f"Meta Description: {page['description']}")
        doc.add_paragraph()

        # ── Hero ──
        h2(doc, "Hero Section")
        if page["hero"].get("eyebrow"):
            body(doc, f"Eyebrow: {page['hero']['eyebrow']}")
        body(doc, f"Headline: {page['hero']['h1']}")
        body(doc, f"Subtitle: {page['hero']['subtitle']}")
        if page["hero"].get("badges"):
            h3(doc, "Key Badges")
            for badge in page["hero"]["badges"]:
                bullet(doc, badge)

        # ── Sections ──
        for sec in page.get("sections", []):
            doc.add_paragraph()
            h2(doc, sec["h2"])
            if sec.get("eyebrow"):
                body(doc, f"Section label: {sec['eyebrow']}")
            if sec.get("body"):
                body(doc, sec["body"])

            # Cards
            if sec.get("cards"):
                for card in sec["cards"]:
                    h3(doc, card["title"])
                    if card.get("tag"):
                        body(doc, f"Tag: {card['tag']}")
                    if card.get("highlight"):
                        body(doc, f"Highlight: {card['highlight']}")
                    body(doc, card["desc"])

            # Steps
            if sec.get("steps"):
                for s in sec["steps"]:
                    h3(doc, f"Step {s['n']}: {s['title']}")
                    body(doc, s["desc"])

            # Timeline items
            if sec.get("timeline"):
                for idx, t in enumerate(sec["timeline"], 1):
                    h3(doc, f"{idx}. {t['label']}")
                    body(doc, t["desc"])

            # Risks
            if sec.get("risks"):
                for r in sec["risks"]:
                    h3(doc, r["label"])
                    body(doc, r["desc"])

            # Eligibility + Documents (two lists side by side in source)
            if sec.get("eligibility"):
                h3(doc, "Eligibility Criteria")
                for item in sec["eligibility"]:
                    bullet(doc, item)
            if sec.get("documents"):
                h3(doc, "Documents Required")
                for item in sec["documents"]:
                    bullet(doc, item)

            # Loan cards (with rate/amount)
            if sec.get("loan_cards"):
                for lc in sec["loan_cards"]:
                    h3(doc, lc["title"])
                    body(doc, f"Interest Rate: {lc['rate']}  |  Loan Amount: {lc['amount']}")
                    body(doc, lc["desc"])

            # Currency pairs
            if sec.get("currency_pairs"):
                for cp in sec["currency_pairs"]:
                    h3(doc, f"{cp['pair']} - {cp['full']}")
                    body(doc, cp["desc"])
                    body(doc, f"Lot Size: {cp['lot']}  |  Tick Size: {cp['tick']}")

            # Table
            if sec.get("table"):
                add_table(doc, sec["table"]["headers"], sec["table"]["rows"])

        # ── FAQ ──
        if page.get("faqs"):
            doc.add_paragraph()
            faq_section(doc, page["faqs"])

        # ── CTA ──
        if page.get("cta"):
            doc.add_paragraph()
            h2(doc, "Call to Action")
            h3(doc, page["cta"]["h2"])
            body(doc, page["cta"]["body"])

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = r"E:\shriram-financial\Shriram_Product_Pages_Content.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")


if __name__ == "__main__":
    build_document()
